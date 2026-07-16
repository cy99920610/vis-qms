"""Best-effort file text extraction, shared by the AI assistant and the
in-content search on the browse page."""
import csv
import io

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

MAX_CONTENT_CHARS = 12000
MAX_PDF_PAGES = 40
SUPPORTED_CONTENT_EXTS = (".pdf", ".docx", ".xlsx", ".csv", ".md", ".txt")


def extract_document_text(document):
    """Returns None for unsupported/unreadable formats, "" if the format is
    supported but no text could be found, or the (possibly truncated)
    extracted text. Skips any I/O for unsupported extensions."""
    name = document.file.name.lower()
    if not name.endswith(SUPPORTED_CONTENT_EXTS):
        return None

    try:
        with document.file.open("rb") as f:
            data = f.read()
    except Exception:
        return None

    try:
        if name.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(data))
            parts = [page.extract_text() or "" for page in reader.pages[:MAX_PDF_PAGES]]
            text = "\n".join(parts)
        elif name.endswith(".docx"):
            doc = DocxDocument(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(parts)
        elif name.endswith(".xlsx"):
            wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
            parts = []
            for ws in wb.worksheets:
                parts.append(f"[Sheet: {ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) for c in row if c is not None]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
        elif name.endswith(".csv"):
            rows = csv.reader(io.StringIO(data.decode("utf-8", errors="ignore")))
            text = "\n".join(" | ".join(row) for row in rows)
        else:  # .md / .txt
            text = data.decode("utf-8", errors="ignore")
    except Exception:
        return None

    text = text.strip()
    if not text:
        return ""
    if len(text) > MAX_CONTENT_CHARS:
        text = text[:MAX_CONTENT_CHARS] + "\n... [truncated]"
    return text
